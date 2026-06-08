# Defines available agent tools and tool execution behavior.
from __future__ import annotations

import json
import logging
import re
from pathlib import Path
from typing import Any, Callable, Dict, Iterator, List, Optional, Tuple

from .base import BaseTool

logger = logging.getLogger("Plant.ReadFileTool")

SUPPORTED_SUFFIXES = (".txt", ".md", ".json", ".pdf", ".docx", ".doc")
CHUNK_SEP = "##"  # chunk_id = "{relative_posix_path}##{index}"


# ========
# Defines has supported files function for this module workflow.
# ========
def has_supported_files(base_dir: Path) -> bool:
    if not base_dir.is_dir():
        return False
    return any(
        p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
        for p in base_dir.rglob("*")
    )


# ========
# Defines rel posix function for this module workflow.
# ========
def rel_posix(path: Path, base: Path) -> str:
    return path.resolve().relative_to(base.resolve()).as_posix()


# ========
# Defines tokenize function for this module workflow.
# ========
def tokenize(q: str) -> List[str]:
    raw = re.findall(r"[\w\u4e00-\u9fff]+", (q or "").lower())
    out: List[str] = []
    for t in raw:
        if len(t) >= 2:
            out.append(t)
        elif len(t) == 1 and "\u4e00" <= t <= "\u9fff":
            out.append(t)
    return out


# ========
# Defines iter chunks function for this module workflow.
# ========
def iter_chunks(
    text: str, max_size: int = 1200, overlap: int = 150
) -> Iterator[Tuple[int, int, str]]:
    max_size = int(max_size)
    overlap = int(overlap)
    if max_size < 1:
        max_size = 1
    overlap = max(0, min(overlap, max_size - 1))
    n = len(text)
    pos = 0
    while pos < n:
        end = min(pos + max_size, n)
        if end < n:
            window = text[pos:end]
            nl = window.rfind("\n")
            if nl >= max_size // 2:
                end = pos + nl + 1
        chunk = text[pos:end]
        yield pos, end, chunk
        if end >= n:
            break
        nxt = end - overlap
        pos = nxt if nxt > pos else end


# ========
# Defines iter chunks shifted function for this module workflow.
# ========
def iter_chunks_shifted(
    full: str, lo: int, hi: int, max_size: int, overlap: int
) -> Iterator[Tuple[int, int, str]]:
    sub = full[lo:hi]
    for ls, le, ck in iter_chunks(sub, max_size, overlap):
        yield lo + ls, lo + le, ck


# ========
# Defines yaml front matter end exclusive function for this module workflow.
# ========
def yaml_front_matter_end_exclusive(text: str) -> int:
    if not text.startswith("---"):
        return 0
    first_nl = text.find("\n")
    if first_nl == -1:
        return 0
    if text[:first_nl].strip() != "---":
        return 0
    i = first_nl + 1
    while i < len(text):
        line_end = text.find("\n", i)
        if line_end == -1:
            line = text[i:]
            nxt = len(text)
        else:
            line = text[i:line_end]
            nxt = line_end + 1
        if line.strip() == "---":
            return nxt
        i = nxt
    return 0


# ========
# Defines heading break offsets function for this module workflow.
# ========
def heading_break_offsets(body: str) -> List[int]:
    lines = body.splitlines(keepends=True)
    breaks: List[int] = []
    pos = 0
    in_fence = False
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        line_start = pos
        s = line.lstrip()
        if s.startswith("```") or s.startswith("~~~"):
            in_fence = not in_fence
            pos += len(line)
            i += 1
            continue
        if not in_fence:
            if re.match(r"^#{1,6}\s+", s):
                breaks.append(line_start)
            elif s.strip() and i + 1 < n:
                s2 = lines[i + 1].lstrip()
                is_eq = bool(re.match(r"^=+\s*$", s2))
                is_dash = bool(re.match(r"^-+\s*$", s2))
                if is_eq:
                    breaks.append(line_start)
                    pos += len(line) + len(lines[i + 1])
                    i += 2
                    continue
                if is_dash:
                    if re.match(r"^\s*[-*+]\s", s) or re.match(r"^\s*\d+\.\s", s):
                        pass
                    else:
                        breaks.append(line_start)
                        pos += len(line) + len(lines[i + 1])
                        i += 2
                        continue
        pos += len(line)
        i += 1
    return breaks


# ========
# Defines line start offsets function for this module workflow.
# ========
def line_start_offsets(text: str) -> List[int]:
    if not text:
        return [0]
    out = [0]
    for i, c in enumerate(text):
        if c == "\n":
            out.append(i + 1)
    return out


# ========
# Defines md heading breaks markdown it function for this module workflow.
# ========
def md_heading_breaks_markdown_it(body: str) -> Optional[List[int]]:
    try:
        from markdown_it import MarkdownIt  # type: ignore[import-not-found]
    except ImportError:
        return None
    try:
        md = MarkdownIt("commonmark")
        tokens = md.parse(body)
    except Exception:
        logger.debug("markdown-it 解析失敗，改用啟發式標題偵測", exc_info=True)
        return None
    line_starts = line_start_offsets(body)
    nlines = len(line_starts)
    breaks: List[int] = []
    for t in tokens:
        if t.type != "heading_open" or not t.map:
            continue
        li = t.map[0]
        if 0 <= li < nlines:
            breaks.append(line_starts[li])
    return sorted(set(breaks))


# ========
# Defines spans from breaks function for this module workflow.
# ========
def spans_from_breaks(
    text: str, breaks: List[int], start_floor: int
) -> List[Tuple[int, int]]:
    br = sorted({b for b in breaks if b >= start_floor})
    if not br:
        return []
    spans: List[Tuple[int, int]] = []
    if br[0] > start_floor:
        spans.append((start_floor, br[0]))
    for j, b in enumerate(br):
        end = br[j + 1] if j + 1 < len(br) else len(text)
        spans.append((b, end))
    return spans


# ========
# Defines paragraph spans local function for this module workflow.
# ========
def paragraph_spans_local(sub: str) -> List[Tuple[int, int]]:
    if not sub.strip():
        return []
    pat = re.compile(r"\n\s*\n+")
    spans: List[Tuple[int, int]] = []
    start = 0
    for m in pat.finditer(sub):
        end = m.start()
        if end > start and sub[start:end].strip():
            spans.append((start, end))
        start = m.end()
    if start < len(sub) and sub[start:].strip():
        spans.append((start, len(sub)))
    if not spans:
        spans.append((0, len(sub)))
    return spans


# ========
# Defines merge spans to budget function for this module workflow.
# ========
def merge_spans_to_budget(
    text: str, spans: List[Tuple[int, int]], max_size: int
) -> List[Tuple[int, int]]:
    if not spans:
        return []
    out: List[Tuple[int, int]] = []
    i = 0
    while i < len(spans):
        j = i
        acc_start, acc_end = spans[i]
        acc_len = acc_end - acc_start
        while j + 1 < len(spans):
            b0, b1 = spans[j + 1]
            gap = text[acc_end:b0]
            add = len(gap) + (b1 - b0)
            if acc_len + add <= max_size:
                j += 1
                acc_end = b1
                acc_len += add
            else:
                break
        out.append((acc_start, acc_end))
        i = j + 1
    return out


# ========
# Defines iter chunks by sections then window function for this module workflow.
# ========
def iter_chunks_by_sections_then_window(
    text: str,
    suffix: str,
    max_size: int,
    overlap: int,
) -> Iterator[Tuple[int, int, str]]:
    if suffix not in (".md", ".txt"):
        yield from iter_chunks(text, max_size, overlap)
        return

    cursor = 0
    if suffix == ".md":
        fm_end = yaml_front_matter_end_exclusive(text)
        if fm_end > 0:
            yield from emit_sized_block(0, text[:fm_end], max_size, overlap)
            cursor = fm_end
    if cursor >= len(text):
        return
    rest = text[cursor:]
    if not rest.strip():
        return

    rel_breaks: List[int]
    if suffix == ".md":
        md_breaks = md_heading_breaks_markdown_it(rest)
        if md_breaks is None:
            rel_breaks = heading_break_offsets(rest)
        else:
            rel_breaks = md_breaks
    else:
        rel_breaks = heading_break_offsets(rest)
    spans: List[Tuple[int, int]]
    if rel_breaks:
        abs_breaks = [cursor + off for off in rel_breaks]
        spans = spans_from_breaks(text, abs_breaks, cursor)
    else:
        raw_paras = [
            (cursor + a, cursor + b) for a, b in paragraph_spans_local(rest)
        ]
        spans = merge_spans_to_budget(text, raw_paras, max_size)

    if not spans:
        yield from iter_chunks_shifted(text, cursor, len(text), max_size, overlap)
        return

    for a, b in spans:
        if a >= b:
            continue
        section = text[a:b]
        if not section.strip():
            continue
        yield from emit_sized_block(a, section, max_size, overlap)


# ========
# Defines emit sized block function for this module workflow.
# ========
def emit_sized_block(
    base: int,
    block: str,
    max_size: int,
    overlap: int,
) -> Iterator[Tuple[int, int, str]]:
    max_size = max(1, int(max_size))
    overlap = max(0, min(int(overlap), max_size - 1))
    if len(block) <= max_size:
        yield base, base + len(block), block
        return
    for ls, le, sub in iter_chunks(block, max_size, overlap):
        yield base + ls, base + le, sub


# ========
# Defines json dumps safe function for this module workflow.
# ========
def json_dumps_safe(data: Any) -> str:
    try:
        return json.dumps(data, ensure_ascii=False, indent=2)
    except (TypeError, ValueError):
        try:
            return json.dumps(
                data, ensure_ascii=False, indent=2, default=str, allow_nan=True
            )
        except (TypeError, ValueError):
            return repr(data)


# ========
# Defines json top level pieces function for this module workflow.
# ========
def json_top_level_pieces(obj: Any) -> List[str]:
    if isinstance(obj, dict):
        keys = list(obj.keys())
        out: List[str] = []
        for k in keys:
            out.append(json_dumps_safe({k: obj[k]}))
        return out
    if isinstance(obj, list):
        return [json_dumps_safe(item) for item in obj]
    return [json_dumps_safe(obj)]


# ========
# Defines iter json structured then window function for this module workflow.
# ========
def iter_json_structured_then_window(
    text: str, max_size: int, overlap: int
) -> Iterator[Tuple[int, int, str]]:
    try:
        obj = json.loads(text)
    except (json.JSONDecodeError, ValueError, TypeError):
        yield from iter_chunks(text, max_size, overlap)
        return
    pieces = json_top_level_pieces(obj)
    if not pieces:
        yield from iter_chunks(text, max_size, overlap)
        return
    sep = "\n\n"
    spans: List[Tuple[int, int]] = []
    pos = 0
    for i, p in enumerate(pieces):
        if i:
            pos += len(sep)
        start = pos
        pos += len(p)
        spans.append((start, pos))
    full = sep.join(pieces)
    i = 0
    n = len(pieces)
    while i < n:
        j = i
        acc = pieces[i]
        acc_start = spans[i][0]
        acc_end = spans[i][1]
        while j + 1 < n and len(acc) + len(sep) + len(pieces[j + 1]) <= max_size:
            j += 1
            acc = acc + sep + pieces[j]
            acc_end = spans[j][1]
        block = full[acc_start:acc_end]
        yield from emit_sized_block(acc_start, block, max_size, overlap)
        i = j + 1


# ========
# Defines iter pdf page chunks function for this module workflow.
# ========
def iter_pdf_page_chunks(
    path: Path, max_size: int, overlap: int
) -> Iterator[Tuple[int, int, str]]:
    import PyPDF2

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    if not pages:
        yield 0, 0, ""
        return
    sep = "\n\n"
    parts: List[str] = []
    for idx, p in enumerate(pages):
        parts.append(p if p else f"[page {idx + 1} empty]")
    full = sep.join(parts)
    pos = 0
    for i, p in enumerate(parts):
        if i:
            pos += len(sep)
        start = pos
        end = pos + len(p)
        pos = end
        block = full[start:end]
        yield from emit_sized_block(start, block, max_size, overlap)


# ========
# Defines iter docx paragraph chunks function for this module workflow.
# ========
def iter_docx_paragraph_chunks(
    path: Path, max_size: int, overlap: int
) -> Iterator[Tuple[int, int, str]]:
    from docx import Document

    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if not paras:
        yield 0, 0, ""
        return
    sep = "\n\n"
    spans: List[Tuple[int, int]] = []
    pos = 0
    for i, t in enumerate(paras):
        if i:
            pos += len(sep)
        start = pos
        pos += len(t)
        spans.append((start, pos))
    full = sep.join(paras)
    i = 0
    n = len(paras)
    while i < n:
        j = i
        acc = paras[i]
        acc_start = spans[i][0]
        acc_end = spans[i][1]
        while j + 1 < n and len(acc) + len(sep) + len(paras[j + 1]) <= max_size:
            j += 1
            acc = acc + sep + paras[j]
            acc_end = spans[j][1]
        block = full[acc_start:acc_end]
        yield from emit_sized_block(acc_start, block, max_size, overlap)
        i = j + 1


# ========
# Defines iter index chunks function for this module workflow.
# ========
def iter_index_chunks(
    *,
    path: Path,
    suffix: str,
    text: str,
    max_size: int,
    overlap: int,
    read_text: Callable[[Path, str], str],
) -> Iterator[Tuple[int, int, str]]:
    if suffix == ".json":
        yield from iter_json_structured_then_window(text, max_size, overlap)
        return
    if suffix == ".pdf":
        try:
            yield from iter_pdf_page_chunks(path, max_size, overlap)
        except Exception:
            logger.warning("PDF 結構切塊失敗，改為全文窗口：%s", path, exc_info=True)
            yield from iter_chunks(read_text(path, suffix), max_size, overlap)
        return
    if suffix in (".docx", ".doc"):
        try:
            yield from iter_docx_paragraph_chunks(path, max_size, overlap)
        except Exception:
            logger.warning("DOCX 結構切塊失敗，改為全文窗口：%s", path, exc_info=True)
            yield from iter_chunks(read_text(path, suffix), max_size, overlap)
        return
    yield from iter_chunks_by_sections_then_window(text, suffix, max_size, overlap)


# ========
# Defines score chunk function for this module workflow.
# ========
def score_chunk(query_tokens: List[str], chunk_lower: str) -> float:
    if not query_tokens:
        return 0.0
    s = 0.0
    for tok in query_tokens:
        s += chunk_lower.count(tok)
    return s


# ========
# Defines ReadFileTool class for this module workflow.
# ========
class ReadFileTool(BaseTool):
    name = "read_file"
    description = (
        "讀取專案 doc/ 目錄參考檔（.txt, .md, .json, .pdf, .docx）。"
        "索引：.md 依 markdown-it-py 標題斷點（失敗則啟發式）、.txt 依啟發式標題／段落；"
        ".json 依頂層結構；.pdf 依頁；.docx 依段落；過長再細切。"
        "必須明確填 action。建議流程：action=search_chunks 用 query 與 top_k 檢索相關片段 → "
        "再用 action=read_chunks 帶 chunk_ids 讀全文片段 → 最後由你綜合（Synthesize）。"
        "只有已知檔名且檔案不大時，才用 action=read_full 一次讀整檔。"
    )
    parameters = {
        "action": {
            "type": "string",
            "description": "必填。read_full | search_chunks | read_chunks",
            "required": True,
        },
        "file_path": {
            "type": "string",
            "description": "相對於 doc/ 的路徑；read_full 必填，其餘 action 可不填",
            "required": False,
        },
        "query": {
            "type": "string",
            "description": "search_chunks 時必填：檢索關鍵字或問句",
            "required": False,
        },
        "chunk_ids": {
            "type": "string",
            "description": "read_chunks 時必填：chunk id，多個以英文逗號分隔（來自 search_chunks 結果）",
            "required": False,
        },
        "top_k": {
            "type": "integer",
            "description": "search_chunks 必填。回傳筆數上限，必須是大於 0 的整數",
            "required": False,
        },
        "output_format": {
            "type": "string",
            "description": "僅 read_full：text 或 json_summary（預設 text）",
            "required": False,
        },
    }

    # Defines __init__ function for this module workflow.
    def __init__(
        self,
        base_dir: Optional[Path] = None,
        *,
        chunk_max_chars: Optional[int] = None,
        chunk_overlap: int = 0,
    ):
        self.base_dir = Path(base_dir) if base_dir else Path("doc")
        cm = None if chunk_max_chars is None else max(1, int(chunk_max_chars))
        co = max(0, int(chunk_overlap))
        co = min(co, cm - 1) if cm and cm > 1 else 0
        self.chunk_max_chars = cm
        self.chunk_overlap = co
        self.file_sig: Dict[str, Tuple[float, int]] = {}
        self.chunks: List[Dict[str, Any]] = []
        self.chunk_by_id: Dict[str, Dict[str, Any]] = {}

    # Defines reset session function for this module workflow.
    def reset_session(self) -> None:
        self.file_sig.clear()
        self.chunks.clear()
        self.chunk_by_id.clear()

    # Defines safe resolve function for this module workflow.
    def safe_resolve(self, file_path: str) -> Tuple[Optional[Path], Optional[str]]:
        try:
            path = (self.base_dir / file_path.strip()).resolve()
            base_resolved = self.base_dir.resolve()
            path.relative_to(base_resolved)
        except ValueError:
            return None, "錯誤：不允許讀取 doc 目錄以外的檔案。"
        except Exception as e:
            return None, f"錯誤：路徑無效：{e}"
        if not path.is_file():
            return None, f"錯誤：檔案不存在或非檔案：{path}"
        return path, None

    # Defines read text by type function for this module workflow.
    def read_text_by_type(self, path: Path, suffix: str) -> str:
        if suffix in (".txt", ".md", ".json"):
            return path.read_text(encoding="utf-8", errors="replace")
        if suffix == ".pdf":
            import PyPDF2

            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
        if suffix in (".docx", ".doc"):
            from docx import Document

            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        raise ValueError(
            f"不支援的副檔名 {suffix}，僅支援 .txt, .md, .json, .pdf, .docx。"
        )

    # Defines current file signature function for this module workflow.
    def current_file_signature(self) -> Dict[str, Tuple[float, int]]:
        sig: Dict[str, Tuple[float, int]] = {}
        if not self.base_dir.is_dir():
            return sig
        for p in self.base_dir.rglob("*"):
            if not p.is_file():
                continue
            if p.suffix.lower() not in SUPPORTED_SUFFIXES:
                continue
            try:
                rel = rel_posix(p, self.base_dir)
                st = p.stat()
                sig[rel] = (st.st_mtime, st.st_size)
            except (OSError, ValueError):
                continue
        return sig

    # Defines rebuild index function for this module workflow.
    def rebuild_index(self) -> None:
        self.chunks.clear()
        self.chunk_by_id.clear()
        if not self.base_dir.is_dir():
            self.file_sig = {}
            return
        for rel, _ in sorted(self.file_sig.items()):
            path = self.base_dir / rel
            if not path.is_file():
                continue
            suffix = path.suffix.lower()
            try:
                text = self.read_text_by_type(path, suffix)
            except ImportError as e:
                logger.warning("索引略過 %s：%s", rel, e)
                continue
            except Exception as e:
                logger.warning("索引讀取失敗 %s: %s", rel, e)
                continue
            i = 0
            try:
                for start, end, chunk_text in iter_index_chunks(
                    path=path,
                    suffix=suffix,
                    text=text,
                    max_size=self.chunk_max_chars or max(1, len(text)),
                    overlap=self.chunk_overlap,
                    read_text=self.read_text_by_type,
                ):
                    if not chunk_text.strip():
                        continue
                    cid = f"{rel}{CHUNK_SEP}{i}"
                    preview = chunk_text.strip().replace("\n", " ")
                    row = {
                        "chunk_id": cid,
                        "file_path": rel,
                        "chunk_index": i,
                        "char_start": start,
                        "char_end": end,
                        "text": chunk_text,
                        "preview": preview,
                    }
                    self.chunks.append(row)
                    self.chunk_by_id[cid] = row
                    i += 1
            except Exception as e:
                logger.warning("索引切塊失敗 %s: %s", rel, e, exc_info=True)
                continue

    # Defines ensure index function for this module workflow.
    def ensure_index(self) -> Optional[str]:
        sig = self.current_file_signature()
        if sig == self.file_sig and self.chunks:
            return None
        self.file_sig = sig
        self.rebuild_index()
        return None

    # Defines execute function for this module workflow.
    def execute(self, **kwargs) -> str:
        raw_action = kwargs.get("action")
        if not isinstance(raw_action, str) or not raw_action.strip():
            return "錯誤：read_file 請明確提供 action：read_full、search_chunks 或 read_chunks。"
        action = raw_action.strip().lower()

        if action == "search_chunks":
            return self.execute_search_chunks(kwargs)
        if action == "read_chunks":
            return self.execute_read_chunks(kwargs)
        if action == "read_full":
            return self.execute_read_full(kwargs)
        return f"錯誤：不支援的 action「{action}」，請用 read_full、search_chunks 或 read_chunks。"

    # Defines execute read full function for this module workflow.
    def execute_read_full(self, kwargs: Dict[str, Any]) -> str:
        file_path = kwargs.get("file_path")
        output_format = (kwargs.get("output_format") or "text").strip()
        if not file_path or not isinstance(file_path, str):
            return "錯誤：read_full 請提供 file_path。"
        if output_format not in ("text", "json_summary"):
            return "錯誤：output_format 僅支援 text 或 json_summary。"

        path, err = self.safe_resolve(file_path)
        if err:
            return err
        assert path is not None
        suffix = path.suffix.lower()
        try:
            text = self.read_text_by_type(path, suffix)
        except ImportError as e:
            return f"錯誤：缺少依賴（{e}），無法讀取 {suffix} 檔案。"
        except Exception as e:
            logger.warning("read_file 讀取失敗 %s: %s", path, e)
            return f"錯誤：無法讀取檔案：{e}"

        if output_format == "text":
            return text
        payload = {
            "file_path": str(path.relative_to(self.base_dir.resolve())),
            "suffix": suffix,
            "char_count": len(text),
            "preview": text,
        }
        return json.dumps(payload, ensure_ascii=False)

    # Defines execute search chunks function for this module workflow.
    def execute_search_chunks(self, kwargs: Dict[str, Any]) -> str:
        query = kwargs.get("query")
        if not query or not isinstance(query, str) or not query.strip():
            return "錯誤：search_chunks 請提供非空的 query。"
        self.ensure_index()
        if not self.chunks:
            return "（索引為空：doc/ 下無可索引之支援檔案，或檔案讀取失敗。）"

        top_k = kwargs.get("top_k")
        if top_k is not None and isinstance(top_k, int) and top_k >= 1:
            k = top_k
        else:
            return "錯誤：search_chunks 請提供 top_k，且必須是大於 0 的整數。"

        q_tokens = tokenize(query)
        ranked: List[Tuple[float, Dict[str, Any]]] = []
        for row in self.chunks:
            txt = row.get("text") or ""
            if not isinstance(txt, str):
                txt = str(txt)
            low = txt.lower()
            sc = score_chunk(q_tokens, low)
            if sc > 0:
                ranked.append((sc, row))
        ranked.sort(key=lambda x: -x[0])
        top = ranked[:k]

        out = []
        for sc, row in top:
            out.append(
                {
                    "chunk_id": row["chunk_id"],
                    "file_path": row["file_path"],
                    "score": round(sc, 4),
                    "preview": row["preview"],
                }
            )
        if not out:
            return json.dumps(
                {
                    "matches": [],
                    "hint": "無命中 chunk。可換關鍵字、縮短 query，或改用 read_full 讀已知檔名。",
                },
                ensure_ascii=False,
            )

        return json.dumps(
            {
                "matches": out,
                "hint": "下一步請對需要的項目使用 action=read_chunks 並帶入 chunk_ids（逗號分隔），再自行綜合（Synthesize）。",
            },
            ensure_ascii=False,
        )

    # Defines execute read chunks function for this module workflow.
    def execute_read_chunks(self, kwargs: Dict[str, Any]) -> str:
        raw = kwargs.get("chunk_ids")
        ids: List[str] = []
        if isinstance(raw, list):
            ids = [str(x).strip() for x in raw if str(x).strip()]
        elif isinstance(raw, str) and raw.strip():
            ids = [x.strip() for x in raw.split(",") if x.strip()]
        if not ids:
            return "錯誤：read_chunks 請提供 chunk_ids（逗號分隔字串）。"

        self.ensure_index()
        parts: List[str] = []
        missing: List[str] = []
        for cid in ids:
            row = self.chunk_by_id.get(cid)
            if not row:
                missing.append(cid)
                continue
            block = (
                f"=== chunk_id: {cid} | file: {row['file_path']} | "
                f"bytes [{row['char_start']}:{row['char_end']}] ===\n"
                f"{row['text']}\n"
            )
            parts.append(block)

        if not parts and missing:
            return (
                "錯誤：找不到任何有效的 chunk_id。請先 search_chunks 取得正確 id，"
                f"無效 id：{missing}"
            )

        header = (
            "【read_chunks】以下為請求的片段全文，請據此與其他工具結果自行綜合（Synthesize）。\n\n"
        )
        body = "".join(parts)
        if missing:
            body += f"\n（以下 id 不存在或索引已更新：{missing}）\n"
        return header + body
